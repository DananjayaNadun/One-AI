"""Attachment pipeline tests. Run: python -m tests.test_attachments"""
import io
import os
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

os.environ["ONEAI_SKIP_DOTENV"] = "1"   # ignore the developer's .env
os.environ["DEFAULT_MODE"] = "panel"    # these suites exercise the panel path

TMP = Path(tempfile.mkdtemp())
os.environ["DB_PATH"] = str(TMP / "att.db")
os.environ["OPENROUTER_KEY"] = "sk-or-test"
os.environ["SECRET_KEY"] = "test"

import attachments as att_lib  # noqa: E402
import config  # noqa: E402
import db  # noqa: E402
import llm  # noqa: E402
import models  # noqa: E402
import app as app_module  # noqa: E402

PASS, FAIL = [], []


def check(name, cond, detail=""):
    (PASS if cond else FAIL).append(name)
    print(f"{'PASS' if cond else 'FAIL'}  {name}{'  -> ' + detail if detail and not cond else ''}")


# --- fixtures ---------------------------------------------------------------

def make_png(w=2400, h=1200, mode="RGBA"):
    from PIL import Image, ImageDraw
    img = Image.new(mode, (w, h), (30, 60, 180, 255) if mode == "RGBA" else (30, 60, 180))
    ImageDraw.Draw(img).rectangle([50, 50, w - 50, h - 50], fill=(240, 200, 40, 255) if mode == "RGBA" else (240, 200, 40))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def make_pdf(text="Reliability score model for COD orders.", pages=2):
    from pypdf import PdfWriter
    import zlib
    # Build a minimal PDF with a real text stream.
    def page_obj(n):
        content = f"BT /F1 14 Tf 72 700 Td ({text} Page {n}) Tj ET".encode()
        return content
    from pypdf.generic import DecodedStreamObject, NameObject, DictionaryObject, ArrayObject, NumberObject, TextStringObject
    writer = PdfWriter()
    for i in range(pages):
        writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def make_real_pdf(text="COD reliability scoring for Sri Lankan sellers.", pages=2):
    """A PDF with genuinely extractable text, built by hand."""
    objects = []
    kids = []
    contents_ids = []
    obj_id = 4
    for p in range(pages):
        stream = f"BT /F1 16 Tf 72 700 Td ({text} Page {p + 1}) Tj ET".encode("latin-1")
        contents_ids.append(obj_id + 1)
        objects.append((obj_id, f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
                                f"/Resources << /Font << /F1 3 0 R >> >> /Contents {obj_id + 1} 0 R >>".encode()))
        objects.append((obj_id + 1, b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"))
        kids.append(f"{obj_id} 0 R")
        obj_id += 2

    header = b"%PDF-1.4\n"
    body = []
    body.append((1, b"<< /Type /Catalog /Pages 2 0 R >>"))
    body.append((2, f"<< /Type /Pages /Kids [{' '.join(kids)}] /Count {pages} >>".encode()))
    body.append((3, b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"))
    body.extend(objects)
    body.sort()

    out = io.BytesIO()
    out.write(header)
    offsets = {}
    for num, content in body:
        offsets[num] = out.tell()
        out.write(f"{num} 0 obj\n".encode() + content + b"\nendobj\n")
    xref_at = out.tell()
    count = max(offsets) + 1
    out.write(f"xref\n0 {count}\n".encode())
    out.write(b"0000000000 65535 f \n")
    for n in range(1, count):
        out.write(f"{offsets.get(n, 0):010d} 00000 n \n".encode())
    out.write(f"trailer\n<< /Size {count} /Root 1 0 R >>\nstartxref\n{xref_at}\n%%EOF".encode())
    return out.getvalue()


def make_docx(paragraphs=("OrderFlow spec", "COD reliability score drives the ranking.")):
    import docx
    document = docx.Document()
    for p in paragraphs:
        document.add_paragraph(p)
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# --- tests ------------------------------------------------------------------

def main():
    application = app_module.app
    application.config["TESTING"] = True
    client = application.test_client()

    # ---------- ingestion ----------
    png = make_png()
    att = att_lib.ingest("diagram.png", png)
    check("image classified", att.kind == "image")
    check("image converted to jpeg data url", att.data_url.startswith("data:image/jpeg;base64,"))
    check("oversized image downscaled",
          att.meta["width"] <= config.MAX_IMAGE_DIMENSION and att.meta["height"] <= config.MAX_IMAGE_DIMENSION,
          f"{att.meta['width']}x{att.meta['height']}")
    check("original dimensions recorded", att.meta["original_size"] == "2400x1200", att.meta["original_size"])
    import base64
    decoded = base64.b64decode(att.data_url.split(",", 1)[1])
    check("image payload under limit", len(decoded) <= config.MAX_IMAGE_BYTES,
          f"{len(decoded)} bytes")
    check("RGBA flattened without crash", len(decoded) > 0)

    pdf = make_real_pdf()
    att = att_lib.ingest("spec.pdf", pdf)
    check("pdf classified as document", att.kind == "document")
    check("pdf text extracted", "reliability" in att.text.lower(), att.text[:80])
    check("pdf page markers present", "[Page 1]" in att.text and "[Page 2]" in att.text)
    check("pdf page count recorded", att.meta.get("pages") == 2, str(att.meta))

    docx_bytes = make_docx()
    att = att_lib.ingest("spec.docx", docx_bytes)
    check("docx text extracted", "OrderFlow" in att.text, att.text[:60])
    check("docx table rows included", "Metric | Value" in att.text, att.text[-60:])

    att = att_lib.ingest("app.py", b"def main():\n    return 'hello'\n")
    check("code file read", "def main" in att.text and att.kind == "text")

    att = att_lib.ingest("data.csv", "name,total\nNimal,LKR 4500\n".encode("utf-16"))
    check("utf-16 decoded", "Nimal" in att.text, att.text[:40])

    # ---------- rejections ----------
    def rejects(name, data, fragment):
        try:
            att_lib.ingest(name, data)
            return False, "no error raised"
        except att_lib.AttachmentError as exc:
            return fragment.lower() in str(exc).lower(), str(exc)

    ok, msg = rejects("virus.exe", b"MZ\x90\x00", "not accepted")
    check("executable rejected", ok, msg)
    ok, msg = rejects("thing.xyz", b"data", "unsupported")
    check("unknown extension rejected", ok, msg)
    ok, msg = rejects("empty.txt", b"", "empty")
    check("empty file rejected", ok, msg)
    ok, msg = rejects("huge.txt", b"x" * (config.MAX_UPLOAD_BYTES + 1), "too large")
    check("oversized file rejected", ok, msg)
    ok, msg = rejects("fake.png", b"this is not an image", "not a readable image")
    check("fake image rejected", ok, msg)
    ok, msg = rejects("scan.pdf", make_real_pdf(text="")[:200], "could not read")
    check("corrupt pdf rejected", ok, msg)

    # path traversal in the filename must not survive
    att = att_lib.ingest("../../etc/passwd.txt", b"root:x:0:0")
    check("path traversal stripped from name", att.name == "passwd.txt", att.name)

    # truncation
    long_text = ("line\n" * 200_000).encode()
    att = att_lib.ingest("big.txt", long_text)
    check("long text truncated", att.truncated and len(att.text) == config.MAX_EXTRACTED_CHARS)

    # descriptor must not leak the full text
    d = att.descriptor()
    check("descriptor omits full text", "text" not in d and len(d["preview"]) <= 181, str(list(d)))

    # ---------- upload endpoint ----------
    data = {"files": [
        (io.BytesIO(make_png(400, 300)), "shot.png"),
        (io.BytesIO(make_real_pdf()), "spec.pdf"),
        (io.BytesIO(b"print(1)"), "main.py"),
    ]}
    r = client.post("/api/upload", data=data, content_type="multipart/form-data")
    body = r.get_json()
    check("upload accepts mixed batch", r.status_code == 200 and len(body["attachments"]) == 3,
          str(r.status_code) + " " + str(body)[:120])
    ids = [a["id"] for a in body["attachments"]]
    check("image descriptor has thumbnail",
          any(a["kind"] == "image" and a["thumbnail"].startswith("data:") for a in body["attachments"]))

    # partial failure: one good, one bad
    r = client.post("/api/upload", data={"files": [
        (io.BytesIO(b"print(1)"), "ok.py"),
        (io.BytesIO(b"MZ"), "bad.exe"),
    ]}, content_type="multipart/form-data")
    body = r.get_json()
    check("partial upload keeps good file",
          r.status_code == 200 and len(body["attachments"]) == 1 and len(body["rejected"]) == 1,
          str(body)[:140])

    r = client.post("/api/upload", data={"files": [(io.BytesIO(b"MZ"), "bad.exe")]},
                    content_type="multipart/form-data")
    check("all-rejected returns 415", r.status_code == 415, str(r.status_code))

    r = client.post("/api/upload", data={}, content_type="multipart/form-data")
    check("no files returns 400", r.status_code == 400)

    r = client.post("/api/upload", data={"files": [
        (io.BytesIO(b"x"), f"f{i}.txt") for i in range(config.MAX_FILES_PER_MESSAGE + 2)
    ]}, content_type="multipart/form-data")
    check("too many files rejected", r.status_code == 400, str(r.status_code))

    # ---------- chat with attachments ----------
    captured = {}

    def fake_generate(prompt, history=None, attachments=None):
        captured["prompt"] = prompt
        captured["attachments"] = attachments or []
        return llm.PanelResult(answer="Reviewed.", nodes={"logic": "a"})

    llm.generate = fake_generate

    r = client.post("/api/chat", json={"message": "What is this?", "chat_id": None,
                                       "attachment_ids": ids})
    body = r.get_json()
    check("chat accepts attachments", r.status_code == 200, str(body)[:140])
    check("attachments reached the model", len(captured["attachments"]) == 3,
          str(len(captured.get("attachments", []))))
    check("response echoes attachments", len(body["attachments"]) == 3)
    chat_id = body["chat_id"]

    # empty message with a file is allowed
    r2 = client.post("/api/upload", data={"files": [(io.BytesIO(b"hi"), "n.txt")]},
                     content_type="multipart/form-data")
    lone = r2.get_json()["attachments"][0]["id"]
    r = client.post("/api/chat", json={"message": "  ", "chat_id": chat_id,
                                       "attachment_ids": [lone]})
    check("file with no message allowed", r.status_code == 200, str(r.get_json())[:120])
    check("default prompt substituted", "review" in captured["prompt"].lower(), captured["prompt"])

    # empty message with NO file is still rejected
    check("empty message without file rejected",
          client.post("/api/chat", json={"message": "  ", "chat_id": chat_id}).status_code == 400)

    # bad ids
    check("malformed attachment id rejected",
          client.post("/api/chat", json={"message": "hi", "attachment_ids": ["../x"]}).status_code == 400)
    check("unknown attachment id rejected",
          client.post("/api/chat", json={"message": "hi", "attachment_ids": ["a" * 32]}).status_code == 400)
    check("non-list attachment_ids rejected",
          client.post("/api/chat", json={"message": "hi", "attachment_ids": "abc"}).status_code == 400)

    # ---------- persistence ----------
    msgs = client.get(f"/api/chats/{chat_id}").get_json()
    first_user = next(m for m in msgs if m["role"] == "user")
    check("attachments persisted on message", len(first_user["attachments"]) == 3,
          str(len(first_user["attachments"])))
    check("stored image keeps data url",
          any(a["kind"] == "image" and a["data_url"].startswith("data:") for a in first_user["attachments"]))

    # cascade: deleting the chat removes attachments
    client.delete(f"/api/chats/{chat_id}")
    with db.get_conn() as conn:
        left = conn.execute("SELECT COUNT(*) c FROM attachments WHERE id IN "
                            f"({','.join('?' * len(ids))})", ids).fetchone()["c"]
    check("attachments cascade-deleted with chat", left == 0, f"{left} remain")

    # ---------- multimodal message construction ----------
    text_only = llm.build_user_content("hello", [])
    check("no attachments -> plain string", isinstance(text_only, str))

    doc_att = {"kind": "document", "name": "spec.pdf", "text": "CONTENT", "truncated": False}
    out = llm.build_user_content("summarise", [doc_att])
    check("document inlined as string", isinstance(out, str) and "CONTENT" in out and "spec.pdf" in out)

    img_att = {"kind": "image", "name": "shot.png", "data_url": "data:image/jpeg;base64,AAA"}
    out = llm.build_user_content("what is this", [img_att])
    check("image -> content parts list", isinstance(out, list) and len(out) == 2, str(type(out)))
    check("text part first", out[0]["type"] == "text")
    check("image part carries data url", out[1]["image_url"]["url"].startswith("data:image/jpeg"))

    out = llm.build_user_content("both", [img_att, doc_att])
    check("mixed batch keeps doc text in the text part",
          "CONTENT" in out[0]["text"] and out[1]["type"] == "image_url")

    check("has_images detects image", llm.has_images([img_att]) and not llm.has_images([doc_att]))

    # aggregator turn must retain the image
    turn = llm._aggregator_turn(out, "PANEL NOTES")
    check("aggregator turn keeps image part",
          isinstance(turn["content"], list)
          and any(p["type"] == "image_url" for p in turn["content"])
          and "PANEL NOTES" in turn["content"][0]["text"])

    turn = llm._aggregator_turn("just text", "PANEL NOTES")
    check("aggregator turn handles plain string",
          isinstance(turn["content"], str) and "PANEL NOTES" in turn["content"])

    # ---------- vision routing ----------
    models._cache.update({"fetched_at": 9e18, "models": [
        {"id": "vendor/text-only:free", "pricing": {"prompt": "0", "completion": "0"},
         "context_length": 100000, "architecture": {"input_modalities": ["text"]}},
        {"id": "vendor/sees:free", "pricing": {"prompt": "0", "completion": "0"},
         "context_length": 200000, "architecture": {"input_modalities": ["text", "image"]}},
        {"id": "vendor/paid-vision", "pricing": {"prompt": "0.005", "completion": "0.01"},
         "context_length": 900000, "architecture": {"input_modalities": ["text", "image"]}},
    ]})
    check("discovery finds free vision model", models.discover(vision=True) == ["vendor/sees:free"],
          str(models.discover(vision=True)))
    check("discovery excludes paid models", "vendor/paid-vision" not in models.discover(),
          str(models.discover()))
    check("text discovery excludes nothing free", "vendor/text-only:free" in models.discover(),
          str(models.discover()))
    check("legacy modality string supported",
          models._supports_images({"architecture": {"modality": "text+image->text"}}))

    print(f"\n{len(PASS)} passed, {len(FAIL)} failed")
    if FAIL:
        print("Failures:", ", ".join(FAIL))
    return 1 if FAIL else 0


if __name__ == "__main__":
    sys.exit(main())
