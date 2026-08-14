"""Attachment ingestion.

Three kinds of file are handled differently:

- images    -> resized, re-encoded, and sent to a vision model as a data URL
- documents -> text extracted server-side (PDF, DOCX) and inlined into the prompt
- text/code -> decoded and inlined directly

Extraction happens on the server, not in the browser, because the browser cannot
read PDF or DOCX and because we must not trust a client-supplied file type.
"""
from __future__ import annotations

import base64
import io
import logging
import mimetypes
import uuid
from dataclasses import dataclass, field
from pathlib import Path

import config

log = logging.getLogger(__name__)

TEXT_SUFFIXES = {
    ".txt", ".md", ".markdown", ".rst", ".log", ".csv", ".tsv", ".json", ".jsonl",
    ".yaml", ".yml", ".xml", ".toml", ".ini", ".cfg", ".env", ".sql",
    ".py", ".js", ".jsx", ".ts", ".tsx", ".html", ".htm", ".css", ".scss",
    ".java", ".c", ".h", ".cpp", ".hpp", ".cs", ".go", ".rs", ".rb", ".php",
    ".swift", ".kt", ".sh", ".bash", ".ps1", ".dart", ".r", ".m", ".pl", ".lua",
}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}
PDF_SUFFIXES = {".pdf"}
DOCX_SUFFIXES = {".docx"}

# Extensions we refuse outright. Nothing here is executed by the server, but
# accepting them invites the user to think the app will run their code.
BLOCKED_SUFFIXES = {".exe", ".dll", ".so", ".bat", ".cmd", ".msi", ".apk", ".jar"}


class AttachmentError(ValueError):
    """A file we cannot or will not process."""


@dataclass
class Attachment:
    id: str
    name: str
    kind: str                    # "image" | "document" | "text"
    size: int
    text: str = ""               # extracted text, for document/text kinds
    data_url: str = ""           # base64 data URL, for image kind
    truncated: bool = False
    meta: dict = field(default_factory=dict)

    def descriptor(self) -> dict:
        """What the browser needs; never includes the full extracted text."""
        return {
            "id": self.id,
            "name": self.name,
            "kind": self.kind,
            "size": self.size,
            "truncated": self.truncated,
            "preview": (self.text[:180] + "\u2026") if len(self.text) > 180 else self.text,
            "thumbnail": self.data_url if self.kind == "image" else "",
            **self.meta,
        }


def _suffix(name: str) -> str:
    return Path(name).suffix.lower()


def classify(name: str) -> str:
    suffix = _suffix(name)
    if suffix in BLOCKED_SUFFIXES:
        raise AttachmentError(f"{suffix} files are not accepted.")
    if suffix in IMAGE_SUFFIXES:
        return "image"
    if suffix in PDF_SUFFIXES or suffix in DOCX_SUFFIXES:
        return "document"
    if suffix in TEXT_SUFFIXES:
        return "text"
    raise AttachmentError(
        f"Unsupported file type '{suffix or name}'. "
        "Supported: images, PDF, DOCX, and text or code files."
    )


# --- Extractors --------------------------------------------------------------

def _decode_text(raw: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_pdf(raw: bytes) -> tuple[str, dict]:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise AttachmentError("PDF support requires 'pypdf'. Run: pip install pypdf")

    try:
        reader = PdfReader(io.BytesIO(raw))
    except Exception as exc:
        raise AttachmentError(f"Could not read that PDF: {exc}")

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:
            raise AttachmentError("That PDF is password protected.")

    pages = []
    for index, page in enumerate(reader.pages[: config.MAX_PDF_PAGES], start=1):
        try:
            content = (page.extract_text() or "").strip()
        except Exception:
            content = ""
        if content:
            pages.append(f"[Page {index}]\n{content}")

    if not pages:
        raise AttachmentError(
            "No text found in that PDF -- it is probably a scan. "
            "Export it as an image and attach that instead, so a vision model can read it."
        )

    meta = {"pages": len(reader.pages)}
    return "\n\n".join(pages), meta


def _extract_docx(raw: bytes) -> tuple[str, dict]:
    try:
        import docx
    except ImportError:
        raise AttachmentError("DOCX support requires 'python-docx'. Run: pip install python-docx")

    try:
        document = docx.Document(io.BytesIO(raw))
    except Exception as exc:
        raise AttachmentError(f"Could not read that document: {exc}")

    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    if not parts:
        raise AttachmentError("That document appears to be empty.")
    return "\n".join(parts), {"paragraphs": len(parts)}


def _prepare_image(raw: bytes, name: str) -> tuple[str, dict]:
    """Re-encode and downscale so we send a sane payload to the model."""
    try:
        from PIL import Image, ImageOps
    except ImportError:
        raise AttachmentError("Image support requires 'Pillow'. Run: pip install Pillow")

    try:
        image = Image.open(io.BytesIO(raw))
        image.load()
    except Exception:
        raise AttachmentError("That file is not a readable image.")

    original = f"{image.width}x{image.height}"

    # Honour EXIF orientation, otherwise phone photos arrive rotated.
    image = ImageOps.exif_transpose(image)

    if image.mode in ("RGBA", "LA", "P"):
        background = Image.new("RGB", image.size, (255, 255, 255))
        converted = image.convert("RGBA")
        background.paste(converted, mask=converted.split()[-1])
        image = background
    elif image.mode != "RGB":
        image = image.convert("RGB")

    longest = max(image.size)
    if longest > config.MAX_IMAGE_DIMENSION:
        scale = config.MAX_IMAGE_DIMENSION / longest
        image = image.resize(
            (max(1, int(image.width * scale)), max(1, int(image.height * scale))),
            Image.LANCZOS,
        )

    buffer = io.BytesIO()
    quality = 85
    image.save(buffer, format="JPEG", quality=quality, optimize=True)
    # Step the quality down rather than sending a payload the API will reject.
    while buffer.tell() > config.MAX_IMAGE_BYTES and quality > 40:
        quality -= 15
        buffer = io.BytesIO()
        image.save(buffer, format="JPEG", quality=quality, optimize=True)

    encoded = base64.b64encode(buffer.getvalue()).decode("ascii")
    meta = {
        "width": image.width,
        "height": image.height,
        "original_size": original,
    }
    return f"data:image/jpeg;base64,{encoded}", meta


# --- Entry point -------------------------------------------------------------

def ingest(filename: str, raw: bytes) -> Attachment:
    """Validate and process one uploaded file."""
    name = Path(filename or "file").name.strip() or "file"
    size = len(raw)

    if size == 0:
        raise AttachmentError(f"'{name}' is empty.")
    if size > config.MAX_UPLOAD_BYTES:
        limit_mb = config.MAX_UPLOAD_BYTES / (1024 * 1024)
        raise AttachmentError(f"'{name}' is too large (limit {limit_mb:.0f} MB).")

    kind = classify(name)
    suffix = _suffix(name)
    attachment = Attachment(id=uuid.uuid4().hex, name=name, kind=kind, size=size)

    if kind == "image":
        attachment.data_url, attachment.meta = _prepare_image(raw, name)
    elif suffix in PDF_SUFFIXES:
        attachment.text, attachment.meta = _extract_pdf(raw)
    elif suffix in DOCX_SUFFIXES:
        attachment.text, attachment.meta = _extract_docx(raw)
    else:
        attachment.text = _decode_text(raw)
        if not attachment.text.strip():
            raise AttachmentError(f"'{name}' contains no readable text.")

    if attachment.text and len(attachment.text) > config.MAX_EXTRACTED_CHARS:
        attachment.text = attachment.text[: config.MAX_EXTRACTED_CHARS]
        attachment.truncated = True

    log.info("ingested %s (%s, %d bytes)", name, kind, size)
    return attachment


def guess_mime(name: str) -> str:
    return mimetypes.guess_type(name)[0] or "application/octet-stream"
